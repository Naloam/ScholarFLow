from __future__ import annotations

from pathlib import Path

try:
    from docx import Document  # type: ignore
except Exception:  # pragma: no cover
    Document = None

import re

from services.workspace import project_root


def _export_dir(project_id: str) -> Path:
    path = project_root(project_id) / "exports"
    path.mkdir(parents=True, exist_ok=True)
    return path


def export_markdown(project_id: str, content: str, export_id: str) -> str:
    path = _export_dir(project_id) / f"{export_id}.md"
    path.write_text(content, encoding="utf-8")
    return str(path)

def _replace_citations(text: str) -> str:
    return re.sub(r"\[(\d+)\]", r"\\cite{ref\1}", text)


def _markdown_table_to_latex(rows: list[list[str]]) -> list[str]:
    if not rows:
        return []
    cols = len(rows[0])
    header = rows[0]
    body = rows[1:]
    latex = []
    latex.append("\\begin{tabular}{" + " | ".join(["l"] * cols) + "}")
    latex.append(" \\hline")
    latex.append(" & ".join(header) + " \\\\ \\hline")
    for row in body:
        latex.append(" & ".join(row) + " \\\\")
    latex.append(" \\hline")
    latex.append("\\end{tabular}")
    return latex


def _parse_table_block(lines: list[str], start: int) -> tuple[list[list[str]], int] | None:
    if start + 1 >= len(lines):
        return None
    if "|" not in lines[start] or "|" not in lines[start + 1]:
        return None
    if not re.search(r"[-:]+", lines[start + 1]):
        return None
    rows: list[list[str]] = []
    idx = start
    while idx < len(lines) and "|" in lines[idx]:
        cells = [c.strip() for c in lines[idx].strip().strip("|").split("|")]
        if idx == start + 1:
            idx += 1
            continue
        rows.append(cells)
        idx += 1
    return rows, idx


def export_latex(project_id: str, content: str, export_id: str) -> str:
    path = _export_dir(project_id) / f"{export_id}.tex"
    raw_lines = content.splitlines()
    lines: list[str] = []
    idx = 0
    while idx < len(raw_lines):
        line = raw_lines[idx]
        table = _parse_table_block(raw_lines, idx)
        if table:
            rows, idx = table
            lines.extend(_markdown_table_to_latex(rows))
            continue
        if line.startswith("# "):
            lines.append(f"\\section{{{line[2:].strip()}}}")
        elif line.startswith("## "):
            lines.append(f"\\subsection{{{line[3:].strip()}}}")
        else:
            lines.append(_replace_citations(line))
        idx += 1
    latex = (
        "\\documentclass{article}\n\\begin{document}\n"
        + "\n".join(lines)
        + "\n\\bibliographystyle{plain}\n\\bibliography{references}\n"
        + "\\end{document}\n"
    )
    path.write_text(latex, encoding="utf-8")
    return str(path)


def export_word(project_id: str, content: str, export_id: str) -> str:
    if Document is None:
        raise RuntimeError("python-docx not available")
    path = _export_dir(project_id) / f"{export_id}.docx"
    doc = Document()
    lines = content.splitlines()
    idx = 0
    while idx < len(lines):
        line = lines[idx]
        table = _parse_table_block(lines, idx)
        if table:
            rows, idx = table
            if rows:
                cols = len(rows[0])
                table_doc = doc.add_table(rows=len(rows), cols=cols)
                for r, row in enumerate(rows):
                    for c, cell in enumerate(row):
                        table_doc.cell(r, c).text = cell
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        else:
            doc.add_paragraph(line)
        idx += 1
    doc.save(str(path))
    return str(path)
